#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

root = Path(__file__).resolve().parents[1]
out_dir = root / "docs"
out_dir.mkdir(parents=True, exist_ok=True)

doc = Document()

# Title
title = doc.add_paragraph()
run = title.add_run("Akasa Air – Data Engineering Solution")
run.font.size = Pt(20)
run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Version: 1.0")
doc.add_paragraph("Generated: Automated documentation")
doc.add_page_break()

# Overview
doc.add_heading("Overview", level=1)
doc.add_paragraph(
    "This document describes a dual-path data engineering pipeline that ingests CSV and XML, "
    "persists data in MySQL (SQLAlchemy ORM), and computes KPIs via both SQL and Pandas."
)

# Architecture
doc.add_heading("Architecture", level=1)
arch = (
    "Data Sources -> Validation/Cleaning -> (A) MySQL via SQLAlchemy -> SQL KPIs -> outputs/*.json\n"
    "                                       (B) Pandas in-memory      -> Pandas KPIs -> outputs/*.csv"
)
p = doc.add_paragraph()
p.add_run(arch).font.name = "Courier New"

# ASCII ER diagram
doc.add_heading("Entity Relationship (ER)", level=2)
ascii_er = (
    "+------------+        +--------+\n"
    "| CUSTOMERS  |        | ORDERS |\n"
    "+------------+        +--------+\n"
    "| customer_id|<---+   |order_id|\n"
    "| name       |    |   |mobile  |\n"
    "| mobile UK  |    +---|date    |\n"
    "| region     |        |sku_id  |\n"
    "+------------+        |amount  |\n"
    "                     +--------+\n"
)
doc.add_paragraph().add_run(ascii_er).font.name = "Courier New"

# Data Flow
doc.add_heading("Data Flow", level=1)
ascii_flow = (
    "CSV(customers)  XML(orders)\n"
    "       |            |\n"
    "   parse/clean  parse/clean\n"
    "       |            |\n"
    "       +-----+------+\n"
    "             |\n"
    "   +---------+----------+\n"
    "   |    src.main        |\n"
    "   +---------+----------+\n"
    "             |\n"
    "     +-------+-------+\n"
    "     |               |\n"
    " MySQL (SQLAlchemy)  Pandas\n"
    "     |               |\n"
    "   SQL KPIs      Pandas KPIs\n"
    "     |               |\n"
    " outputs/*.json  outputs/*.csv\n"
)
doc.add_paragraph().add_run(ascii_flow).font.name = "Courier New"

# KPIs
doc.add_heading("KPIs", level=1)
t = doc.add_table(rows=1, cols=3)
hdr = t.rows[0].cells
hdr[0].text = "KPI"
hdr[1].text = "Description"
hdr[2].text = "Implementation"
rows = [
    ("Repeat Customers", "Customers with >1 order", "SQL GROUP BY/HAVING; Pandas groupby().size()"),
    ("Monthly Order Trends", "Orders and revenue per month", "SQL YEAR/MONTH + SUM; Pandas dt.year/month"),
    ("Regional Revenue", "Revenue, counts by region", "SQL JOIN + SUM/AVG; Pandas merge + agg"),
    ("Top Spenders (N days)", "Top customers by spend in last N days", "SQL date filter + agg; Pandas date filter + agg"),
]
for r in rows:
    row = t.add_row().cells
    row[0].text, row[1].text, row[2].text = r

# Setup & Run
doc.add_heading("Setup & Run", level=1)
doc.add_paragraph("1. Create venv and install requirements: pip install -r requirements.txt")
doc.add_paragraph("2. Configure .env (DB creds, optional DB_SOCKET=/tmp/mysql.sock)")
doc.add_paragraph("3. Run: python -m src.main")

# Configuration
doc.add_heading("Configuration", level=1)
doc.add_paragraph(".env keys: DB_HOST, DB_PORT, DB_USER, DB_PASSWORD, DB_NAME, DB_SOCKET (optional)")

# Troubleshooting
doc.add_heading("Troubleshooting", level=1)
doc.add_paragraph("- Ensure MySQL is running and accessible via UNIX socket or TCP.")
doc.add_paragraph("- Verify data paths: data/customers.csv, data/orders.xml.")

# Save
out_path = out_dir / "AkasaAir-DataEngineering-Solution.docx"
doc.save(out_path)
print(f"Wrote {out_path}")
